from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from dotenv import load_dotenv
from docx import Document
import datetime
import os
import pandas as pd
import pandasql as ps
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import base64

load_dotenv()

SPREADSHEET_ID = os.getenv('FORM_RESPONSES_SHEET')
SENDER_EMAIL = os.getenv('EMAIL')
SCOPES = [
    'https://www.googleapis.com/auth/spreadsheets',
    'https://www.googleapis.com/auth/gmail.send'
]

COLUMN_NAMES = ["submission_id", "srn", "request_type", "timestamp", "status"]

def post_latest_event():
    creds = Credentials.from_authorized_user_file("v3/token.json", SCOPES)
    service = build("sheets", "v4", credentials=creds)
    sheet_name = 'Form responses 1'
    range_name = f'{sheet_name}!A1:K100'
    result = service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=range_name
    ).execute()
    values = result.get("values", [])
    print(f"ead {len(values)} rows from '{sheet_name}'")
    headers = values[0]
    data_rows = values[1:]
    headers
    data_rows
    df = pd.DataFrame(data=data_rows,columns=headers[:len(data_rows[0])])
    latest_row = df.iloc[-1]  # last row
    srn = latest_row.get('SRN', '')
    tor = latest_row.get('The type of request', '')
    time = latest_row.get('Timestamp', '')

    # --- Prepare data for writing ---
    dest_sheet_name = 'procces tracker'
    write_range = f"{dest_sheet_name}!A1"

    row_to_write = [[f"{srn}-{tor}-{time}", srn, tor, time, "created"]]

    body = {"values": row_to_write}

    # --- Append to destination sheet ---
    result = (
        service.spreadsheets()
        .values()
        .append(
            spreadsheetId=SPREADSHEET_ID,
            range=write_range,
            valueInputOption="USER_ENTERED",
            insertDataOption="INSERT_ROWS",
            body=body
        )
        .execute()
    )
    print(result)


def read_and_send():
    creds = Credentials.from_authorized_user_file("v3/token.json", SCOPES)
    service = build("sheets", "v4", credentials=creds)
    sheet_name = 'procces tracker'
    dest_name = 'Form responses 1'
    range_name = f'{sheet_name}!A1:K100'
    result = service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=range_name
    ).execute()
    values = result.get("values", [])
    print(f"ead {len(values)} rows from '{sheet_name}'")
    data_rows = values[0:]
    df = pd.DataFrame(data_rows, columns=COLUMN_NAMES)
    query = """
    SELECT DISTINCT submission_id
    FROM df
    GROUP BY submission_id
    HAVING COUNT(DISTINCT status) = 1
    AND MAX(status) = 'created'
    """
    result = ps.sqldf(query)
    # if result:
    #     for i in result:
    #         print('send email')
    print(result)
    srn,tor,time=result['submission_id'][0].split('-')

    range_name = f'{dest_name}!A1:K100'
    result = service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=range_name
    ).execute()
    values = result.get("values", [])
    print(f"ead {len(values)} rows from '{sheet_name}'")
    headers = values[0]
    data_rows = values[1:]
    headers
    data_rows
    df = pd.DataFrame(data=data_rows,columns=headers[:len(data_rows[0])])
    matched_rows = df[
    (df['SRN'] == srn) &
    (df['The type of request'] == tor) &
    (df['Timestamp'] == time)
    ]   
    doc = Document()
    doc.add_heading("REVA UNIVERSITY", level=1)
    doc.add_paragraph("SCHOOL OF ARCHITECTURE\n")
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y')}\n")
    doc.add_paragraph("Permission Letter for Case Study\n", style='Intense Quote')

    body = (
        f"To whomsoever it may concern\n\n"
        f"The bearer of this letter, {matched_rows['Name'][0]} bearing SRN {matched_rows['SRN'][0]} "
        f"is the bonafide student of School of Architecture, REVA University, "
        f"Bangalore studying {matched_rows['Semester'][0]} semester B.Arch. for the academic year {datetime.datetime.now().strftime('%Y')}.\n\n"
        f"As a part of their curriculum for {matched_rows['Subject'][0]}, he is visiting below mentioned places "
        f"between {matched_rows['from date'][0]} and {matched_rows['to date'][0]}.\n\n"
        f"{matched_rows['Places to visit'][0]}\n\n"
        "Kindly permit him to enter the premises to carry out the detailed case study. "
        "This visit will be focused on procuring site plan, floor plans, study materials and any other data. "
        "If any, capturing photographs and interviews where necessary. We assure you that the material collected "
        "will be used solely for academic purpose. Please consider this as an official request letter. "
        "Your co-operation in this matter will be highly appreciated.\n\n"
        "Yours sincerely,\n\n"
        "Prof. Ar. Sudhir Acharya\n"
        "Director\n"
        "School of Architecture\n"
        "REVA University\n"
        "Bangalore"
    )

    doc.add_paragraph(body)
    filename = f"CaseStudy_Letter_{srn}_{tor}.docx"
    doc.save(filename)
    creds = Credentials.from_authorized_user_file('v3/token.json', SCOPES)
    service = build('gmail', 'v1', credentials=creds)

    # Email details
    subject = "Permission Letter for Case Study - REVA University"
    body_text = (
        "Dear Student,\n\n"
        "Please find attached your permission letter for Case Study.\n\n"
        "Regards,\nREVA University Automation System"
    )

    # Create message
    message = MIMEMultipart()
    message['to'] = SENDER_EMAIL
    message['subject'] = subject
    message.attach(MIMEText(body_text, "plain"))
    with open(filename, "rb") as f:
        attachment = MIMEApplication(f.read(), _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document")
        attachment.add_header("Content-Disposition", "attachment", filename=os.path.basename(filename))
        message.attach(attachment)

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    try:
        send_message = (
            service.users().messages().send(userId="me", body={"raw": raw_message}).execute()
        )
    except Exception as e:
        print(e)

if __name__ == '__main__':
    post_latest_event()
    read_and_send()